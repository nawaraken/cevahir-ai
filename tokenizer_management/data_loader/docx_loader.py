import os
import logging
import traceback
from typing import Optional, List
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

logger = logging.getLogger(__name__)

# ======================================================
#  Özel Hata Sınıfları
# ======================================================

class DOCXLoaderError(Exception):
    pass

class FileFormatError(DOCXLoaderError):
    pass

class DocumentProcessingError(DOCXLoaderError):
    pass

# ======================================================
#  DOCXLoader Sınıfı
# ======================================================

class DOCXLoader:
    """
    DOCXLoader, DOCX dosyalarını yükler, içeriğini çıkarır ve temizler.
    
    Özellikler:
    - Paragrafları ve tabloları işler
    - Chunking ile büyük metin dosyalarını işler
    - Stil özelliklerini çıkarma desteği
    - Boş hücreleri atlama ve kontrol karakterlerini temizleme
    """

    def __init__(self, 
                 max_chunk_size: int = 512, 
                 extract_style: bool = False,
                 ignore_empty_cells: bool = True):
        """
        Args:
            max_chunk_size (int): Çıkarılacak maksimum metin parçası boyutu.
            extract_style (bool): Stil özelliklerini çıkar. Varsayılan: False.
            ignore_empty_cells (bool): Boş hücreleri atla. Varsayılan: True.
        """
        if not isinstance(max_chunk_size, int) or max_chunk_size <= 0:
            raise ValueError(f"`max_chunk_size` pozitif bir tamsayı olmalıdır, ancak {max_chunk_size} alındı.")
        
        self.max_chunk_size = max_chunk_size
        self.extract_style = extract_style
        self.ignore_empty_cells = ignore_empty_cells

    # ======================================================
    #  Dosya Yükleme
    # ======================================================

    def load_file(self, file_path: str) -> dict:
        """
        DOCX dosyasını yükler ve çıkarır.
        
        Args:
            file_path (str): Yüklenecek DOCX dosyasının tam yolu.

        Returns:
            dict: Çıkarılmış ve temizlenmiş metin + meta bilgileri.
        
        Raises:
            FileFormatError: Desteklenmeyen format hatası.
            DocumentProcessingError: Dosya açma veya işleme hatası.
            Exception: Diğer genel hatalar.
        """
        if not file_path.endswith('.docx'):
            raise FileFormatError(f"Geçersiz dosya formatı: {file_path}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")

        logger.info(f" DOCX dosyası yükleniyor: {file_path}")

        try:
            document = Document(file_path)
            logger.info(f" DOCX dosyası başarıyla yüklendi.")

            texts = []

            #  Paragrafları işle
            texts.extend(self._extract_paragraphs(document))

            #  Tabloları işle
            texts.extend(self._extract_tables(document))

            #  Büyük dosyalarda chunking yapısı uygula
            final_text = self._chunk_text(" ".join(texts))

            logger.info(f" Çıkarılan içerik uzunluğu: {len(final_text)} karakter.")

            #  ÇIKTI FORMATINI DÜZENLİYORUZ 
            return {
                "data": final_text,
                "source_file": os.path.basename(file_path),
                "modality": "text"
            }

        except PackageNotFoundError:
            raise FileFormatError(f"{file_path} geçersiz bir DOCX dosyasıdır veya bozuk.")

        except Exception as e:
            logger.error(f" DOCX işlenirken hata oluştu: {e}", exc_info=True)
            raise DocumentProcessingError(f"DOCX işlenirken hata oluştu: {e}")


    # ======================================================
    #  Paragraf İşleme
    # ======================================================

    def _extract_paragraphs(self, document: Document) -> List[str]:
        texts = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text = paragraph.text.strip()

                if self.extract_style:
                    text = self._extract_style_from_paragraph(paragraph)

                texts.append(text)
        return texts

    # ======================================================
    #  Stil Çıkartma
    # ======================================================

    def _extract_style_from_paragraph(self, paragraph) -> str:
        styled_text = ""
        for run in paragraph.runs:
            text = run.text

            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"__{text}__"

            styled_text += text
        return styled_text.strip()

    # ======================================================
    #  Tablo İşleme
    # ======================================================

    def _extract_tables(self, document: Document) -> List[str]:
        texts = []
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip() or not self.ignore_empty_cells:
                        row_text.append(cell.text.strip())
                if row_text:
                    texts.append(" | ".join(row_text))
        return texts

    # ======================================================
    #  Chunk Yapısı
    # ======================================================

    def _chunk_text(self, text: str) -> str:
        """
        Metni maksimum chunk uzunluğuna göre böler.
        
        Args:
            text (str): Bölünecek metin.

        Returns:
            str: Birleştirilmiş metin.
        """
        if not text:
            return ""

        words = text.split()
        chunks = []
        current_chunk = ""

        for word in words:
            if len(current_chunk) + len(word) + 1 > self.max_chunk_size:
                chunks.append(current_chunk)
                current_chunk = word
            else:
                current_chunk += " " + word if current_chunk else word

        if current_chunk:
            chunks.append(current_chunk)

        return " ".join(chunks)

