from pathlib import Path
import pytest
from docx import Document
from tokenizer_management.data_loader.docx_loader import DOCXLoader

def create_sample_docx(file_path: Path):
    doc = Document()
    doc.add_paragraph("This is a test paragraph.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    doc.save(str(file_path))

def test_docx_loader(tmp_path: Path):
    file = tmp_path / "test.docx"
    create_sample_docx(file)
    loader = DOCXLoader()
    output = loader.load_file(str(file))
    # Paragraf ve tablo hücrelerinden gelen metinlerin çıktıda yer alması beklenir.
    assert "This is a test paragraph." in output
    assert "Cell 1" in output
    assert "Cell 2" in output
